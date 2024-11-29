from docx import Document
import pandas as pd

# Creating the document for Word output
doc = Document()
doc.add_heading('Comparación de Amplificadores de Fibra Óptica', level=1)

# Adding introductory text
doc.add_paragraph("La siguiente tabla muestra una comparación entre los amplificadores de fibra óptica Autonics BF4RP, Omron E3X-DA41-S y Optex D3RF-TDP para aplicaciones de detección.")

# Defining the data for the comparison
data = {
    "Característica": [
        "Velocidad de respuesta",
        "Ajuste de sensibilidad",
        "Funciones de temporizador",
        "Indicadores y visualización",
        "Protección contra interferencias",
        "Consumo de corriente",
        "Salidas y compatibilidad con PLC"
    ],
    "Autonics BF4RP": [
        "Máxima de 0.5 ms (modo normal), 0.7 ms (frecuencia 2)",
        "Ajuste automático y remoto mediante botón o entrada externa",
        "Temporizador de retardo de apagado fijo de 40 ms",
        "Indicadores LED (rojo para operación, verde para estabilidad)",
        "Modo de frecuencia diferencial para evitar interferencias",
        "Máximo de 45 mA",
        "Salida NPN/PNP de colector abierto, 30V DC y 100 mA"
    ],
    "Omron E3X-DA41-S": [
        "Hasta 48 µs en modo super alta velocidad, 4 ms en modo de alta resolución",
        "Enseñanza y ajuste manual, control de potencia automático (APC)",
        "Tres temporizadores (retardo de encendido/apagado y pulso único) ajustables de 1 ms a 5 s",
        "Pantalla digital dual (nivel de incidente y umbral)",
        "Comunicación óptica para prevenir interferencias, hasta 10 unidades",
        "40 mA máximo en modo estándar",
        "Salida NPN/PNP de colector abierto, 26.4V DC y 50 mA"
    ],
    "Optex D3RF-TDP": [
        "16 µs en modo 1-HS, hasta 22 µs enlazado",
        "FALUX para estabilidad de emisión y corrección de sensibilidad automática (ASC)",
        "Configurables (retardo de encendido, apagado, pulso único) entre 0.1 ms y 9999 ms",
        "Pantalla de 7 segmentos de 8 dígitos, opción al 100% para monitoreo de luz recibida",
        "Prevención de interferencias cruzadas (crosstalk) para operar hasta 16 unidades en modo ECO",
        "Hasta 39 mA en modo dual de salida, opción ECO reduce consumo en 31%",
        "Salida NPN/PNP de colector abierto, 30V DC y hasta 100 mA; salida dual configurable"
    ]
}

# Adding table to the Word document
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Característica"
hdr_cells[1].text = "Autonics BF4RP"
hdr_cells[2].text = "Omron E3X-DA41-S"
hdr_cells[3].text = "Optex D3RF-TDP"

for i in range(len(data["Característica"])):
    row_cells = table.add_row().cells
    row_cells[0].text = data["Característica"][i]
    row_cells[1].text = data["Autonics BF4RP"][i]
    row_cells[2].text = data["Omron E3X-DA41-S"][i]
    row_cells[3].text = data["Optex D3RF-TDP"][i]

# Save Word document
doc_path = "D:Comparación_Amplificadores_Fibra_Óptica.docx"
doc.save(doc_path)

# Creating DataFrame for Excel output
df = pd.DataFrame(data)

# Save Excel file
excel_path = "D:Comparación_Amplificadores_Fibra_Óptica.xlsx"
df.to_excel(excel_path, index=False)

# Output paths for user download
doc_path, excel_path
