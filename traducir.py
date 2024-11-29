from docx import Document
from deep_translator import GoogleTranslator

# Función para traducir un archivo de Word
def translate_word_file(input_path, output_path, source_lang='de', target_lang='es'):
    # Cargar el documento de Word
    doc = Document(input_path)
    
    # Preparar el traductor
    translator = GoogleTranslator(source=source_lang, target=target_lang)
    
    # Crear un nuevo documento para la versión traducida
    translated_doc = Document()

    # Traducir cada párrafo del documento
    for para in doc.paragraphs:
        translated_text = translator.translate(para.text)
        translated_doc.add_paragraph(translated_text)
    
    # Guardar el documento traducido
    translated_doc.save(output_path)
    print(f"Documento traducido guardado en: {output_path}")

# Ruta del archivo original en alemán y del archivo traducido
input_file = 'Lista_de_piezas.docx'
output_file = 'Lista_de_piezas_traducida.docx'

# Llamar a la función para traducir
translate_word_file(input_file, output_file)
